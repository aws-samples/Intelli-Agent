import logging
import os
import tempfile
import uuid
from pathlib import Path
from typing import List

import mammoth
from docx import Document as pyDocument
from langchain.docstore.document import Document
from langchain_community.document_loaders.base import BaseLoader
from llm_bot_dep.loaders.html import CustomHtmlLoader
from llm_bot_dep.schemas.processing_parameters import ProcessingParameters
from llm_bot_dep.utils.s3_utils import download_file_from_s3
from PIL import Image

logger = logging.getLogger(__name__)


class CustomDocLoader(BaseLoader):
    """Load docx file.

    Args:
        file_path: File path of the docx file.
        s3_uri: S3 URI of the docx file.
    """

    def __init__(
        self,
        file_path: str,
        s3_uri: str,
    ):
        """Initialize with file path."""
        self.file_path = file_path
        self.s3_uri = s3_uri

    def clean_document(self, doc: pyDocument):
        """Clean document including removing header and footer for each page

        Args:
            doc (Document): The document to clean
        """
        # Remove headers and footers
        for section in doc.sections:
            if section.header is not None:
                for paragraph in section.header.paragraphs:
                    paragraph.clear()

            if section.footer is not None:
                for paragraph in section.footer.paragraphs:
                    paragraph.clear()

    def load(self, image_result_bucket_name: str) -> List[Document]:
        """Load from file path."""

        # Create a directory for images if it doesn't exist
        image_dir = "/tmp/doc_images"
        os.makedirs(image_dir, exist_ok=True)

        def _convert_image(image):
            # Generate unique filename for the image
            image_filename = f"{uuid.uuid4()}.jpg"
            image_path = os.path.join(image_dir, image_filename)

            # Convert and save the image
            with image.open() as image_bytes:
                img = Image.open(image_bytes)
                # Convert to RGB if necessary (in case of PNG with transparency)
                if img.mode in ("RGBA", "P"):
                    img = img.convert("RGB")
                img.save(image_path, "JPEG")

            # Return the image path to be used in the HTML
            return {"src": image_path}

        pyDoc = pyDocument(self.file_path)
        self.clean_document(pyDoc)
        pyDoc.save(self.file_path)

        with open(self.file_path, "rb") as docx_file:
            result = mammoth.convert_to_html(
                docx_file,
                convert_image=mammoth.images.img_element(_convert_image),
            )
            html_content = result.value
            loader = CustomHtmlLoader(file_path=self.file_path, s3_uri=self.s3_uri)
            metadata = {"file_path": self.s3_uri, "file_type": "docx"}
            doc = loader.load(image_result_bucket_name, file_content=html_content)
            doc.metadata = metadata

        return doc


def process_doc(processing_params: ProcessingParameters):
    """Process text content and split into documents.
    
    Args:
        processing_params: ProcessingParameters object containing the bucket and key
        
    Returns:
        List of processed documents.
    """
    bucket = processing_params.source_bucket_name
    key = processing_params.source_object_key
    suffix = Path(key).suffix
    
    # Create a temporary file with .docx suffix
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as temp_file:
        local_path = temp_file.name
    
    # Download the file locally
    download_file_from_s3(bucket, key, local_path)
    
    # Use the loader with the local file path
    loader = CustomDocLoader(file_path=local_path, s3_uri=f"s3://{bucket}/{key}")
    doc = loader.load(image_result_bucket_name=processing_params.portal_bucket_name)
    doc_list = [doc]
    
    # Clean up the temporary file
    Path(local_path).unlink(missing_ok=True)

    return doc_list
